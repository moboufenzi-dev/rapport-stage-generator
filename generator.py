from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64
from datetime import datetime
from PIL import Image as PILImage


def hex_to_rgb(hex_color: str) -> RGBColor:
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def decode_base64_image(base64_str: str) -> io.BytesIO:
    """Décode une image base64 et la convertit en PNG pour compatibilité python-docx."""
    if not base64_str:
        raise ValueError("Image base64 vide")
    if ',' in base64_str:
        base64_str = base64_str.split(',')[1]
    decoded = base64.b64decode(base64_str)

    # Convertir en PNG via Pillow pour assurer la compatibilité
    try:
        input_stream = io.BytesIO(decoded)
        pil_image = PILImage.open(input_stream)

        # Convertir en RGB si nécessaire (pour les images RGBA ou autres modes)
        if pil_image.mode in ('RGBA', 'LA', 'P'):
            # Créer un fond blanc pour les images avec transparence
            background = PILImage.new('RGB', pil_image.size, (255, 255, 255))
            if pil_image.mode == 'P':
                pil_image = pil_image.convert('RGBA')
            background.paste(pil_image, mask=pil_image.split()[-1] if pil_image.mode == 'RGBA' else None)
            pil_image = background
        elif pil_image.mode != 'RGB':
            pil_image = pil_image.convert('RGB')

        # Sauvegarder en PNG
        output_stream = io.BytesIO()
        pil_image.save(output_stream, format='PNG')
        output_stream.seek(0)
        return output_stream
    except Exception as e:
        print(f"[DEBUG] Conversion Pillow échouée: {e}, utilisation directe")
        return io.BytesIO(decoded)


def format_date_fr(date_str: str) -> str:
    if not date_str:
        return "[Date]"
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        mois = ["janvier", "février", "mars", "avril", "mai", "juin",
                "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
        return f"{d.day} {mois[d.month-1]} {d.year}"
    except:
        return date_str


def calculate_duration(date_debut: str, date_fin: str) -> str:
    if not date_debut or not date_fin:
        return "[durée]"
    try:
        d1 = datetime.strptime(date_debut, "%Y-%m-%d")
        d2 = datetime.strptime(date_fin, "%Y-%m-%d")
        months = (d2.year - d1.year) * 12 + d2.month - d1.month
        return f"{months} mois" if months != 1 else "1 mois"
    except:
        return "[durée]"


def add_page_number_field(paragraph):
    """Ajoute un champ PAGE pour le numéro de page."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    # Placeholder text
    text_elem = OxmlElement('w:t')
    text_elem.text = "1"

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(text_elem)
    run._r.append(fldChar3)


def create_toc_entry(doc, text, level=1, page=""):
    """Crée une entrée de table des matières avec tabulation et points de suite."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(4)

    # Indentation selon le niveau
    if level == 1:
        paragraph.paragraph_format.left_indent = Cm(0)
    elif level == 2:
        paragraph.paragraph_format.left_indent = Cm(0.75)
    else:
        paragraph.paragraph_format.left_indent = Cm(1.5)

    # Tabulation avec points de suite à 15cm (pour laisser de la place)
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # Texte du titre
    run = paragraph.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(11)
    elif level == 2:
        run.font.size = Pt(10)
    else:
        run.font.size = Pt(10)
        run.italic = True

    # Tabulation + numéro de page
    paragraph.add_run("\t")
    page_run = paragraph.add_run(page)
    page_run.font.size = Pt(11 if level == 1 else 10)
    if level == 1:
        page_run.bold = True


def create_toc(doc, data):
    """Crée une table des matières avec numérotation automatique."""
    page_num = 3

    if data.include_thanks:
        create_toc_entry(doc, "REMERCIEMENTS", 1, str(page_num))
        page_num += 1

    if data.include_abstract:
        create_toc_entry(doc, "RÉSUMÉ", 1, str(page_num))
        page_num += 1

    for chapter_idx, chapter in enumerate(data.chapters, 1):
        # Numéroter les chapitres dans la TOC
        create_toc_entry(doc, f"{chapter_idx}. {chapter.title}", 1, str(page_num))
        for sub_idx, sub in enumerate(chapter.children, 1):
            # Numéroter les sous-chapitres
            create_toc_entry(doc, f"{chapter_idx}.{sub_idx}. {sub.title}", 2, str(page_num))
            # Sous-sous-chapitres si présents
            if hasattr(sub, 'children') and sub.children:
                for subsub_idx, subsub in enumerate(sub.children, 1):
                    create_toc_entry(doc, f"{chapter_idx}.{sub_idx}.{subsub_idx}. {subsub.title}", 3, str(page_num))
        page_num += 1

    if data.include_annexes:
        create_toc_entry(doc, "ANNEXES", 1, str(page_num))


def setup_document_styles(doc, style_config):
    """Configure les styles du document avec indentation professionnelle."""
    styles = doc.styles

    # Normal - texte avec indentation
    normal = styles['Normal']
    normal.font.name = style_config.font_family
    normal.font.size = Pt(style_config.font_size)
    normal.paragraph_format.line_spacing = style_config.line_spacing
    normal.paragraph_format.first_line_indent = Cm(1.25)  # Alinéa première ligne
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1 - Chapitres principaux (pas d'indentation)
    h1 = styles['Heading 1']
    h1.font.name = style_config.font_family
    h1.font.size = Pt(style_config.title1_size)
    h1.font.bold = style_config.title1_bold
    h1.font.color.rgb = hex_to_rgb(style_config.title1_color)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.left_indent = Cm(0)
    h1.paragraph_format.first_line_indent = Cm(0)

    # Heading 2 - Sous-sections (indentation 1 niveau)
    h2 = styles['Heading 2']
    h2.font.name = style_config.font_family
    h2.font.size = Pt(style_config.title2_size)
    h2.font.bold = style_config.title2_bold
    h2.font.color.rgb = hex_to_rgb(style_config.title2_color)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.left_indent = Cm(0.75)
    h2.paragraph_format.first_line_indent = Cm(0)

    # Heading 3 - Sous-sous-sections (indentation 2 niveaux)
    h3 = styles['Heading 3']
    h3.font.name = style_config.font_family
    h3.font.size = Pt(style_config.title3_size)
    h3.font.italic = style_config.title3_italic
    h3.font.bold = False
    h3.font.color.rgb = hex_to_rgb(style_config.title3_color)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.left_indent = Cm(1.5)
    h3.paragraph_format.first_line_indent = Cm(0)


def setup_header_with_logos(section, data):
    """Configure l'en-tête avec logos - symétrique."""
    header = section.header

    # Supprimer le contenu existant
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Créer un tableau symétrique (largeurs fixes) - total 16cm
    tbl = header.add_table(rows=1, cols=3, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # Largeurs symétriques : gauche 5cm | centre 6cm | droite 5cm
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(6)
    tbl.columns[2].width = Cm(5)

    remove_table_borders(tbl)

    row = tbl.rows[0]

    # Logo école (gauche)
    cell_left = row.cells[0]
    para_left = cell_left.paragraphs[0]
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_left.paragraph_format.first_line_indent = Cm(0)
    if data.logos.logo_ecole and len(data.logos.logo_ecole) > 100:
        try:
            img = decode_base64_image(data.logos.logo_ecole)
            run = para_left.add_run()
            run.add_picture(img, height=Cm(1.2))
        except:
            pass

    # Centre (vide)
    cell_center = row.cells[1]
    para_center = cell_center.paragraphs[0]
    para_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_center.paragraph_format.first_line_indent = Cm(0)

    # Logo entreprise (droite)
    cell_right = row.cells[2]
    para_right = cell_right.paragraphs[0]
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_right.paragraph_format.first_line_indent = Cm(0)
    if data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100:
        try:
            img = decode_base64_image(data.logos.logo_entreprise)
            run = para_right.add_run()
            run.add_picture(img, height=Cm(1.2))
        except:
            pass


def setup_footer_with_page_number(section, data):
    """Configure le pied de page : entreprise à gauche, numéro au centre, nom à droite."""
    footer = section.footer

    # Supprimer le contenu existant
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Créer un tableau symétrique pour le footer - total 16cm
    tbl = footer.add_table(rows=1, cols=3, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # Largeurs symétriques
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(6)
    tbl.columns[2].width = Cm(5)

    remove_table_borders(tbl)

    row = tbl.rows[0]

    # Entreprise (gauche)
    cell_left = row.cells[0]
    para_left = cell_left.paragraphs[0]
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_left.paragraph_format.first_line_indent = Cm(0)
    if data.entreprise_nom:
        run = para_left.add_run(data.entreprise_nom)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Numéro de page (centre)
    cell_center = row.cells[1]
    para_center = cell_center.paragraphs[0]
    para_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_center.paragraph_format.first_line_indent = Cm(0)
    if data.page.show_page_number:
        run = para_center.add_run("- ")
        run.font.size = Pt(9)
        add_page_number_field(para_center)
        run = para_center.add_run(" -")
        run.font.size = Pt(9)

    # Nom étudiant (droite)
    cell_right = row.cells[2]
    para_right = cell_right.paragraphs[0]
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_right.paragraph_format.first_line_indent = Cm(0)
    if data.page.show_student_name and data.nom:
        run = para_right.add_run(f"{data.prenom} {data.nom}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)


def set_cell_shading(cell, color_hex):
    """Applique une couleur de fond à une cellule."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def remove_table_borders(table):
    """Supprime toutes les bordures d'un tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def add_centered_paragraph(doc, space_after=0):
    """Crée un paragraphe centré sans indentation (pour page de garde)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    return para


# ════════════════════════════════════════════════════════════════
#                    MODÈLES DE PAGE DE GARDE
# ════════════════════════════════════════════════════════════════

def generate_cover_classique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style classique - NORMES OFFICIELLES (Compilatio/Scribbr)."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ LOGOS EN HAUT (2cm de haut max) ═══
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=3)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(5)
        logo_table.columns[1].width = Cm(6)
        logo_table.columns[2].width = Cm(5)
        remove_table_borders(logo_table)
        row = logo_table.rows[0]

        if has_logo_ecole:
            try:
                para = row.cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = row.cells[2].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    # ═══ ESPACE ═══
    add_centered_paragraph(doc, 25)

    # ═══ TITRE PRINCIPAL ═══
    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # ═══ SUJET DU STAGE ═══
    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # ═══ IMAGE CENTRALE ═══
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 18)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # École et année (12pt)
    school = add_centered_paragraph(doc, 20)
    run = school.add_run(f"{data.ecole or '[Établissement]'} — {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    # ═══ ENTREPRISE (14pt) ═══
    ent_name = add_centered_paragraph(doc, 4)
    run = ent_name.add_run(f"Stage réalisé chez {data.entreprise_nom or '[Entreprise]'}")
    run.font.size = Pt(14)

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    # ═══ DATES (12pt) ═══
    dates = add_centered_paragraph(doc, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    # ═══ TUTEURS (12pt) ═══
    tuteur_table = doc.add_table(rows=1, cols=2)
    tuteur_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tuteur_table.autofit = False
    tuteur_table.columns[0].width = Cm(7)
    tuteur_table.columns[1].width = Cm(7)
    remove_table_borders(tuteur_table)

    cell_left = tuteur_table.rows[0].cells[0]
    p1 = cell_left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    r1 = p1.add_run("Tuteur entreprise\n")
    r1.font.size = Pt(10)
    run1 = p1.add_run(data.tuteur_nom or "[Nom]")
    run1.bold = True
    run1.font.size = Pt(12)

    cell_right = tuteur_table.rows[0].cells[1]
    p2 = cell_right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run("Tuteur académique\n")
    r2.font.size = Pt(10)
    run2 = p2.add_run(data.tuteur_academique_nom or "[Nom]")
    run2.bold = True
    run2.font.size = Pt(12)


def generate_cover_moderne(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style moderne - NORMES OFFICIELLES avec bandeau coloré."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ BANDEAU SUPÉRIEUR COLORÉ ═══
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Cm(16)
    remove_table_borders(header_table)
    header_cell = header_table.rows[0].cells[0]
    set_cell_shading(header_cell, data.style.title1_color.lstrip('#'))

    # Titre dans le bandeau
    p_header = header_cell.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.first_line_indent = Cm(0)
    p_header.paragraph_format.space_before = Pt(18)
    p_header.paragraph_format.space_after = Pt(6)
    run = p_header.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Sujet dans le bandeau
    if data.sujet_stage:
        p_sujet = add_cell_paragraph(header_cell, 12)
        p_sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.italic = True

    # ═══ LOGOS (2cm) ═══
    add_centered_paragraph(doc, 15)
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(7)
        logo_table.columns[1].width = Cm(7)
        remove_table_borders(logo_table)
        row = logo_table.rows[0]

        if has_logo_ecole:
            try:
                para = row.cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = row.cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    # ═══ IMAGE CENTRALE ═══
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 18)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # École (12pt)
    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  •  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    # ═══ ENTREPRISE (14pt) ═══
    ent_name = add_centered_paragraph(doc, 4)
    run = ent_name.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    # Dates (12pt)
    dates = add_centered_paragraph(doc, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    # ═══ TUTEURS (12pt) ═══
    tuteur_table = doc.add_table(rows=1, cols=2)
    tuteur_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tuteur_table.autofit = False
    tuteur_table.columns[0].width = Cm(7)
    tuteur_table.columns[1].width = Cm(7)
    remove_table_borders(tuteur_table)

    for i, (label, name) in enumerate([
        ("Tuteur entreprise", data.tuteur_nom),
        ("Tuteur académique", data.tuteur_academique_nom)
    ]):
        cell = tuteur_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r1 = p.add_run(f"{label}\n")
        r1.font.size = Pt(10)
        r2 = p.add_run(name or "[Nom]")
        r2.bold = True
        r2.font.size = Pt(12)


def add_cell_paragraph(cell, space_after=0):
    """Crée un paragraphe sans indentation dans une cellule."""
    para = cell.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def generate_cover_elegant(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style élégant - NORMES OFFICIELLES avec ligne verticale."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ STRUCTURE : Ligne verticale + Contenu ═══
    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    main_table.columns[0].width = Cm(0.5)
    main_table.columns[1].width = Cm(15.5)

    stripe = main_table.rows[0].cells[0]
    content = main_table.rows[0].cells[1]

    # Bande colorée
    set_cell_shading(stripe, data.style.title1_color.lstrip('#'))

    # ═══ CONTENU PRINCIPAL ═══
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Logos en haut (2cm)
    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    # Espace
    add_cell_paragraph(content, 25)

    # RAPPORT DE STAGE
    title = add_cell_paragraph(content, 4)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 12)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Image centrale
    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_cell_paragraph(content, 18)
    student = add_cell_paragraph(content, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = add_cell_paragraph(content, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # École (12pt)
    school = add_cell_paragraph(content, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  |  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    # ═══ ENTREPRISE (14pt) ═══
    ent = add_cell_paragraph(content, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_cell_paragraph(content, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    # Dates (12pt)
    dates = add_cell_paragraph(content, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    # ═══ TUTEURS (12pt) ═══
    tuteurs = add_cell_paragraph(content, 4)
    run = tuteurs.add_run(f"Tuteur entreprise : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(12)

    if data.tuteur_academique_nom:
        tut2 = add_cell_paragraph(content, 0)
        run = tut2.add_run(f"Tuteur académique : {data.tuteur_academique_nom}")
        run.font.size = Pt(12)


def generate_cover_minimaliste(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style minimaliste - Ultra épuré, espace blanc dominant."""

    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ GRAND ESPACE BLANC EN HAUT ═══
    for _ in range(5):
        add_centered_paragraph(doc, 18)

    # ═══ TITRE MINIMAL ═══
    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # ═══ LIGNE DÉCORATIVE ═══
    line = add_centered_paragraph(doc, 15)
    run = line.add_run("─────────────")
    run.font.size = Pt(12)
    run.font.color.rgb = primary_color

    # ═══ SUJET ═══
    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 25)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # ═══ IMAGE (optionnel) ═══
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 25)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = add_centered_paragraph(doc, 15)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)

    # ═══ INFOS EN BAS ═══
    add_centered_paragraph(doc, 30)

    # Entreprise
    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Dates
    dates = add_centered_paragraph(doc, 4)
    run = dates.add_run(f"{date_debut_fr} — {date_fin_fr}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def generate_cover_academique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style académique - Cadre double traditionnel."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ CADRE EXTÉRIEUR ═══
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_table.autofit = False
    outer_table.columns[0].width = Cm(16)

    # Bordure extérieure (épaisse)
    set_table_border(outer_table, data.style.title1_color, '24')

    outer_cell = outer_table.rows[0].cells[0]
    add_cell_paragraph(outer_cell, 4)

    # ═══ CADRE INTÉRIEUR ═══
    inner_table = outer_cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inner_table.autofit = False
    inner_table.columns[0].width = Cm(15)

    # Bordure intérieure
    set_table_border(inner_table, data.style.title1_color, '12')

    content = inner_table.rows[0].cells[0]
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # ═══ LOGOS EN HAUT ═══
    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    # ═══ ÉTABLISSEMENT ═══
    add_cell_paragraph(content, 10)
    school_p = add_cell_paragraph(content, 6)
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school_p.add_run(data.ecole or "[Établissement]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    # Formation
    form_p = add_cell_paragraph(content, 15)
    form_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = form_p.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # ═══ TITRE ═══
    title = add_cell_paragraph(content, 6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 12)
        sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Image centrale
    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 15)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_cell_paragraph(content, 10)
    student = add_cell_paragraph(content, 4)
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run("Présenté par")
    run.font.size = Pt(10)

    name_p = add_cell_paragraph(content, 8)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # ═══ ENTREPRISE ═══
    ent = add_cell_paragraph(content, 4)
    ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ent.add_run(f"Stage effectué chez {data.entreprise_nom or '[Entreprise]'}")
    run.font.size = Pt(12)

    # Dates
    dates = add_cell_paragraph(content, 15)
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(11)

    # ═══ TUTEURS ═══
    tuteurs = add_cell_paragraph(content, 4)
    tuteurs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tuteurs.add_run(f"Tuteur entreprise : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)

    if data.tuteur_academique_nom:
        tut2 = add_cell_paragraph(content, 4)
        tut2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tut2.add_run(f"Tuteur académique : {data.tuteur_academique_nom}")
        run.font.size = Pt(10)

    # Année
    add_cell_paragraph(content, 10)
    year_p = add_cell_paragraph(content, 6)
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_p.add_run(data.annee_scolaire or "[Année scolaire]")
    run.font.size = Pt(12)
    run.font.color.rgb = primary_color


def generate_cover_geometrique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style géométrique - Formes modernes et dynamiques."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ BLOC COLORÉ EN HAUT À DROITE (simulé avec tableau) ═══
    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    top_table.columns[0].width = Cm(10)
    top_table.columns[1].width = Cm(6)
    remove_table_borders(top_table)

    # Cellule droite colorée
    right_cell = top_table.rows[0].cells[1]
    set_cell_shading(right_cell, data.style.title1_color.lstrip('#'))

    # Logo école dans la cellule gauche
    left_cell = top_table.rows[0].cells[0]
    if has_logo_ecole:
        try:
            p = left_cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            img = decode_base64_image(data.logos.logo_ecole)
            p.add_run().add_picture(img, height=Cm(2))
        except:
            pass

    # Année dans le bloc coloré
    p_year = right_cell.paragraphs[0]
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.first_line_indent = Cm(0)
    p_year.paragraph_format.space_before = Pt(15)
    p_year.paragraph_format.space_after = Pt(15)
    run = p_year.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.bold = True

    # ═══ ESPACE ═══
    add_centered_paragraph(doc, 20)

    # ═══ TITRE ═══
    title = doc.add_paragraph()
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = doc.add_paragraph()
        sujet.paragraph_format.first_line_indent = Cm(0)
        sujet.paragraph_format.space_after = Pt(15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # ═══ LIGNE DÉCORATIVE ═══
    line_table = doc.add_table(rows=1, cols=1)
    line_table.autofit = False
    line_table.columns[0].width = Cm(5)
    line_cell = line_table.rows[0].cells[0]
    set_cell_shading(line_cell, data.style.title1_color.lstrip('#'))
    remove_table_borders(line_table)
    add_cell_paragraph(line_cell, 0)

    # ═══ IMAGE ═══
    if has_image_centrale:
        try:
            add_centered_paragraph(doc, 15)
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 20)
    student = doc.add_paragraph()
    student.paragraph_format.first_line_indent = Cm(0)
    student.paragraph_format.space_after = Pt(4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = doc.add_paragraph()
    formation.paragraph_format.first_line_indent = Cm(0)
    formation.paragraph_format.space_after = Pt(4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # École
    school = doc.add_paragraph()
    school.paragraph_format.first_line_indent = Cm(0)
    school.paragraph_format.space_after = Pt(20)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ═══ BLOC INFOS EN BAS ═══
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.autofit = False
    info_table.columns[0].width = Cm(8)
    info_table.columns[1].width = Cm(8)
    remove_table_borders(info_table)

    # Entreprise
    p1 = info_table.rows[0].cells[0].paragraphs[0]
    p1.paragraph_format.first_line_indent = Cm(0)
    r1 = p1.add_run("Entreprise : ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(100, 100, 100)
    r2 = p1.add_run(data.entreprise_nom or "[Entreprise]")
    r2.font.size = Pt(11)
    r2.bold = True

    # Dates
    p2 = info_table.rows[0].cells[1].paragraphs[0]
    p2.paragraph_format.first_line_indent = Cm(0)
    r = p2.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    r.font.size = Pt(10)

    # Tuteur
    p3 = info_table.rows[1].cells[0].paragraphs[0]
    p3.paragraph_format.first_line_indent = Cm(0)
    r = p3.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    r.font.size = Pt(10)

    # Logo entreprise en bas à droite
    if has_logo_entreprise:
        try:
            p4 = info_table.rows[1].cells[1].paragraphs[0]
            p4.paragraph_format.first_line_indent = Cm(0)
            p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            img = decode_base64_image(data.logos.logo_entreprise)
            p4.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass


def generate_cover_bicolore(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style bicolore - Division verticale en deux zones."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ STRUCTURE BICOLORE ═══
    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    # 40% coloré | 60% blanc
    main_table.columns[0].width = Cm(6.4)
    main_table.columns[1].width = Cm(9.6)

    left_col = main_table.rows[0].cells[0]
    right_col = main_table.rows[0].cells[1]

    # Colonne gauche colorée
    set_cell_shading(left_col, data.style.title1_color.lstrip('#'))

    # ═══ CONTENU COLONNE GAUCHE (logos + année) ═══
    left_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Logo école
    if has_logo_ecole:
        try:
            add_cell_paragraph(left_col, 15)
            logo_p = add_cell_paragraph(left_col, 15)
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.logo_ecole)
            logo_p.add_run().add_picture(img, height=Cm(2))
        except:
            pass

    # Espace
    for _ in range(3):
        add_cell_paragraph(left_col, 20)

    # "STAGE" en vertical
    stage_p = add_cell_paragraph(left_col, 30)
    stage_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stage_p.add_run("STAGE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Année
    year_p = add_cell_paragraph(left_col, 10)
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_p.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 200, 200)

    # Dates
    dates_p = add_cell_paragraph(left_col, 30)
    dates_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates_p.add_run(f"{date_debut_fr}\n—\n{date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(200, 200, 200)

    # Logo entreprise en bas
    if has_logo_entreprise:
        try:
            for _ in range(2):
                add_cell_paragraph(left_col, 20)
            logo2_p = add_cell_paragraph(left_col, 10)
            logo2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.logo_entreprise)
            logo2_p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass

    # ═══ CONTENU COLONNE DROITE ═══
    right_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Espace en haut
    add_cell_paragraph(right_col, 30)

    # Titre
    title = add_cell_paragraph(right_col, 6)
    run = title.add_run("RAPPORT\nDE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        add_cell_paragraph(right_col, 12)
        sujet = add_cell_paragraph(right_col, 12)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Image
    if has_image_centrale:
        try:
            add_cell_paragraph(right_col, 8)
            img_para = add_cell_paragraph(right_col, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    # Étudiant
    add_cell_paragraph(right_col, 25)
    student = add_cell_paragraph(right_col, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation
    formation = add_cell_paragraph(right_col, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    # École
    school = add_cell_paragraph(right_col, 20)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Entreprise
    ent = add_cell_paragraph(right_col, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_cell_paragraph(right_col, 15)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(11)

    # Tuteurs
    add_cell_paragraph(right_col, 20)
    tut = add_cell_paragraph(right_col, 4)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_pro(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style Pro - Corporate business, sobre et professionnel."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ HEADER BAR ═══
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Cm(5)
    header_table.columns[1].width = Cm(6)
    header_table.columns[2].width = Cm(5)
    remove_table_borders(header_table)

    # Fond coloré pour tout le header
    for cell in header_table.rows[0].cells:
        set_cell_shading(cell, data.style.title1_color.lstrip('#'))

    row = header_table.rows[0]

    # Logo école
    if has_logo_ecole:
        try:
            p = row.cells[0].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            img = decode_base64_image(data.logos.logo_ecole)
            p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass
    else:
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(15)

    # Centre : Année
    p_center = row.cells[1].paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.first_line_indent = Cm(0)
    p_center.paragraph_format.space_before = Pt(15)
    p_center.paragraph_format.space_after = Pt(15)
    run = p_center.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Logo entreprise
    if has_logo_entreprise:
        try:
            p = row.cells[2].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            img = decode_base64_image(data.logos.logo_entreprise)
            p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass

    # ═══ ESPACE ═══
    add_centered_paragraph(doc, 30)

    # ═══ TITRE ═══
    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Image
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 25)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    # Formation + École
    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(f"{data.formation or '[Formation]'}  •  {data.ecole or '[Établissement]'}")
    run.font.size = Pt(12)

    # ═══ BLOC ENTREPRISE ═══
    add_centered_paragraph(doc, 30)

    # Cadre info entreprise
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(8)
    remove_table_borders(info_table)

    # Entreprise
    p1_label = info_table.rows[0].cells[0].paragraphs[0]
    p1_label.paragraph_format.first_line_indent = Cm(0)
    r = p1_label.add_run("Entreprise")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p1_value = info_table.rows[0].cells[1].paragraphs[0]
    p1_value.paragraph_format.first_line_indent = Cm(0)
    r = p1_value.add_run(data.entreprise_nom or "[Entreprise]")
    r.font.size = Pt(12)
    r.bold = True

    # Période
    p2_label = info_table.rows[1].cells[0].paragraphs[0]
    p2_label.paragraph_format.first_line_indent = Cm(0)
    r = p2_label.add_run("Période")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p2_value = info_table.rows[1].cells[1].paragraphs[0]
    p2_value.paragraph_format.first_line_indent = Cm(0)
    r = p2_value.add_run(f"{date_debut_fr} — {date_fin_fr}")
    r.font.size = Pt(11)

    # Tuteur entreprise
    p3_label = info_table.rows[2].cells[0].paragraphs[0]
    p3_label.paragraph_format.first_line_indent = Cm(0)
    r = p3_label.add_run("Tuteur")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p3_value = info_table.rows[2].cells[1].paragraphs[0]
    p3_value.paragraph_format.first_line_indent = Cm(0)
    r = p3_value.add_run(data.tuteur_nom or "[Nom]")
    r.font.size = Pt(11)

    # Tuteur académique
    if data.tuteur_academique_nom:
        p4_label = info_table.rows[3].cells[0].paragraphs[0]
        p4_label.paragraph_format.first_line_indent = Cm(0)
        r = p4_label.add_run("Suivi")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)

        p4_value = info_table.rows[3].cells[1].paragraphs[0]
        p4_value.paragraph_format.first_line_indent = Cm(0)
        r = p4_value.add_run(data.tuteur_academique_nom)
        r.font.size = Pt(11)

    # ═══ FOOTER BAR ═══
    add_centered_paragraph(doc, 30)
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    footer_table.columns[0].width = Cm(16)

    # Ligne colorée en bas
    set_table_border_top(footer_table, data.style.title1_color, '18')
    remove_table_borders(footer_table)

    p_footer = footer_table.rows[0].cells[0].paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.first_line_indent = Cm(0)
    p_footer.paragraph_format.space_before = Pt(10)

    if data.entreprise_ville:
        run = p_footer.add_run(f"{data.entreprise_nom} — {data.entreprise_ville}")
    else:
        run = p_footer.add_run(data.entreprise_nom or "")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def set_table_border(table, color_hex, size='6'):
    """Applique une bordure colorée à un tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color_hex.lstrip('#'))
        tbl_borders.append(border)
    # Intérieurs vides
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_table_border_top(table, color_hex, size='6'):
    """Applique une bordure uniquement en haut du tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    border = OxmlElement('w:top')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:color'), color_hex.lstrip('#'))
    tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def generate_cover_gradient(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style gradient - Dégradé coloré moderne."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ BANDEAU DÉGRADÉ EN HAUT ═══
    gradient_table = doc.add_table(rows=1, cols=1)
    gradient_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gradient_table.autofit = False
    gradient_table.columns[0].width = Cm(16)
    remove_table_borders(gradient_table)

    grad_cell = gradient_table.rows[0].cells[0]
    set_cell_shading(grad_cell, data.style.title1_color.lstrip('#'))

    # Contenu dans le bandeau
    p_banner = grad_cell.paragraphs[0]
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.first_line_indent = Cm(0)
    p_banner.paragraph_format.space_before = Pt(20)
    p_banner.paragraph_format.space_after = Pt(8)

    run = p_banner.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Sujet dans le bandeau
    if data.sujet_stage:
        p_sujet = add_cell_paragraph(grad_cell, 15)
        p_sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.italic = True

    # ═══ LOGOS ═══
    add_centered_paragraph(doc, 15)
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(7)
        logo_table.columns[1].width = Cm(7)
        remove_table_borders(logo_table)

        if has_logo_ecole:
            try:
                para = logo_table.rows[0].cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = logo_table.rows[0].cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    # ═══ IMAGE CENTRALE ═══
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 10)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ ÉTUDIANT ═══
    add_centered_paragraph(doc, 20)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  •  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ═══ ENTREPRISE ═══
    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    dates = add_centered_paragraph(doc, 10)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(11)

    # ═══ TUTEUR ═══
    tut = add_centered_paragraph(doc, 4)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_timeline(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style timeline - Frise chronologique verticale."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ STRUCTURE : Ligne timeline + Contenu ═══
    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    main_table.columns[0].width = Cm(2)
    main_table.columns[1].width = Cm(14)

    timeline_col = main_table.rows[0].cells[0]
    content_col = main_table.rows[0].cells[1]

    # ═══ COLONNE TIMELINE ═══
    timeline_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Point de début (date début)
    p1 = add_cell_paragraph(timeline_col, 4)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("●")
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    date_p1 = add_cell_paragraph(timeline_col, 20)
    date_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p1.add_run(date_debut_fr.split()[0] if date_debut_fr else "")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Ligne verticale (simulée avec |)
    for _ in range(6):
        line_p = add_cell_paragraph(timeline_col, 0)
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line_p.add_run("│")
        run.font.size = Pt(10)
        run.font.color.rgb = primary_color

    # Point de fin (date fin)
    add_cell_paragraph(timeline_col, 20)
    p2 = add_cell_paragraph(timeline_col, 4)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("●")
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    date_p2 = add_cell_paragraph(timeline_col, 0)
    date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p2.add_run(date_fin_fr.split()[0] if date_fin_fr else "")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ═══ COLONNE CONTENU ═══
    content_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Logos
    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content_col.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(6)
        logo_tbl.columns[1].width = Cm(6)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

    # Titre
    add_cell_paragraph(content_col, 15)
    title = add_cell_paragraph(content_col, 4)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = add_cell_paragraph(content_col, 10)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.italic = True

    # Image
    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content_col, 10)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    # Étudiant
    add_cell_paragraph(content_col, 15)
    student = add_cell_paragraph(content_col, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(content_col, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)

    school = add_cell_paragraph(content_col, 10)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Entreprise
    ent = add_cell_paragraph(content_col, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    # Tuteur
    add_cell_paragraph(content_col, 10)
    tut = add_cell_paragraph(content_col, 0)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_creative(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style creative - Design original avec cercles."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # ═══ LOGOS EN HAUT ═══
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(8)
        logo_table.columns[1].width = Cm(8)
        remove_table_borders(logo_table)

        if has_logo_ecole:
            try:
                para = logo_table.rows[0].cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = logo_table.rows[0].cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    # ═══ ESPACE ═══
    add_centered_paragraph(doc, 30)

    # ═══ TITRE STYLISÉ ═══
    # "RAPPORT" en gros
    rapport = add_centered_paragraph(doc, 0)
    run = rapport.add_run("RAPPORT")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = primary_color

    # "DE STAGE" plus petit
    de_stage = add_centered_paragraph(doc, 8)
    run = de_stage.add_run("DE STAGE")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # Ligne décorative colorée
    line = add_centered_paragraph(doc, 15)
    run = line.add_run("━━━━━━━━━━━━━━━━")
    run.font.size = Pt(10)
    run.font.color.rgb = primary_color

    # ═══ SUJET ═══
    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(f"« {data.sujet_stage} »")
        run.font.size = Pt(13)
        run.italic = True

    # ═══ IMAGE ═══
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # ═══ CERCLE INFO (simulé avec tableau arrondi) ═══
    add_centered_paragraph(doc, 20)

    # Étudiant
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  |  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(10)

    # ═══ INFO ENTREPRISE ═══
    # Ligne séparatrice
    sep = add_centered_paragraph(doc, 10)
    run = sep.add_run("─────────────")
    run.font.color.rgb = RGBColor(200, 200, 200)

    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 4)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(10)

    dates = add_centered_paragraph(doc, 8)
    run = dates.add_run(f"{date_debut_fr}  —  {date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Tuteur
    tut = add_centered_paragraph(doc, 4)
    run = tut.add_run(f"Encadré par {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)


def generate_cover_luxe(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style luxe - Élégant avec bordures dorées."""

    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100

    # Couleur dorée
    gold_color = RGBColor(184, 134, 11)

    # ═══ CADRE EXTÉRIEUR DORÉ ═══
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_table.autofit = False
    outer_table.columns[0].width = Cm(16)
    set_table_border(outer_table, 'b8860b', '36')  # 4.5pt border

    outer_cell = outer_table.rows[0].cells[0]

    # ═══ CADRE INTÉRIEUR ═══
    add_cell_paragraph(outer_cell, 6)
    inner_table = outer_cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inner_table.autofit = False
    inner_table.columns[0].width = Cm(15)
    set_table_border(inner_table, 'b8860b', '18')  # 2.25pt border

    content = inner_table.rows[0].cells[0]
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # ═══ LOGOS ═══
    if has_logo_ecole or has_logo_entreprise:
        add_cell_paragraph(content, 6)
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

    # ═══ ÉCOLE ═══
    add_cell_paragraph(content, 10)
    school_p = add_cell_paragraph(content, 4)
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school_p.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(11)
    run.font.color.rgb = gold_color
    run.font.small_caps = True

    # ═══ ORNEMENT ═══
    orn1 = add_cell_paragraph(content, 6)
    orn1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = orn1.add_run("— ✦ —")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color

    # ═══ TITRE ═══
    title = add_cell_paragraph(content, 6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(50, 50, 50)

    # ═══ SUJET ═══
    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 8)
        sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    # ═══ IMAGE ═══
    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 10)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    # ═══ LIGNE DORÉE ═══
    add_cell_paragraph(content, 10)
    line = add_cell_paragraph(content, 8)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("───────────────────")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color

    # ═══ ÉTUDIANT ═══
    student = add_cell_paragraph(content, 6)
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(content, 2)
    formation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)
    run.italic = True

    # ═══ ENTREPRISE ═══
    add_cell_paragraph(content, 10)
    ent = add_cell_paragraph(content, 2)
    ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.font.size = Pt(13)
    run.font.color.rgb = gold_color

    # ═══ DATES ═══
    dates = add_cell_paragraph(content, 4)
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates.add_run(f"{date_debut_fr}  —  {date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ═══ ANNÉE ═══
    add_cell_paragraph(content, 8)
    year = add_cell_paragraph(content, 6)
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run(data.annee_scolaire or "[Année scolaire]")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color
    run.font.small_caps = True


def generate_report(data) -> io.BytesIO:
    """Génère le rapport de stage."""
    doc = Document()

    # Configuration de la page (A4)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(data.page.margin_top)
    section.bottom_margin = Cm(data.page.margin_bottom)
    section.left_margin = Cm(data.page.margin_left)
    section.right_margin = Cm(data.page.margin_right)

    # Première page différente (pas de header/footer sur page de garde)
    section.different_first_page_header_footer = True

    # Configurer les styles
    setup_document_styles(doc, data.style)

    # Variables
    duree = calculate_duration(data.date_debut, data.date_fin)
    date_debut_fr = format_date_fr(data.date_debut)
    date_fin_fr = format_date_fr(data.date_fin)

    # ════════════════════════════════════════════════════════════════
    #                         PAGE DE GARDE
    # ════════════════════════════════════════════════════════════════
    if data.include_cover:
        # Sélection du modèle de page de garde
        cover_model = getattr(data, 'cover_model', 'classique')

        if cover_model == 'moderne':
            generate_cover_moderne(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'elegant':
            generate_cover_elegant(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'minimaliste':
            generate_cover_minimaliste(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'academique':
            generate_cover_academique(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'geometrique':
            generate_cover_geometrique(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'bicolore':
            generate_cover_bicolore(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'pro':
            generate_cover_pro(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'gradient':
            generate_cover_gradient(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'timeline':
            generate_cover_timeline(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'creative':
            generate_cover_creative(doc, data, date_debut_fr, date_fin_fr, duree)
        elif cover_model == 'luxe':
            generate_cover_luxe(doc, data, date_debut_fr, date_fin_fr, duree)
        else:
            generate_cover_classique(doc, data, date_debut_fr, date_fin_fr, duree)

        doc.add_page_break()

    # Configurer header et footer pour les pages suivantes
    setup_header_with_logos(section, data)
    setup_footer_with_page_number(section, data)

    # ════════════════════════════════════════════════════════════════
    #                      TABLE DES MATIÈRES
    # ════════════════════════════════════════════════════════════════
    if data.include_toc:
        toc_heading = doc.add_heading("TABLE DES MATIÈRES", level=1)
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        create_toc(doc, data)

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #                        REMERCIEMENTS
    # ════════════════════════════════════════════════════════════════
    if data.include_thanks:
        thanks_heading = doc.add_heading("REMERCIEMENTS", level=1)
        thanks_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p1 = doc.add_paragraph()
        p1.add_run(f"Je tiens à remercier {data.entreprise_nom or '[Entreprise]'} pour m'avoir accueilli durant ce stage.")

        if data.tuteur_nom:
            p2 = doc.add_paragraph()
            p2.add_run(f"Je remercie particulièrement {data.tuteur_nom}")
            if data.tuteur_poste:
                p2.add_run(f", {data.tuteur_poste},")
            p2.add_run(" pour son encadrement tout au long de ce stage.")

        if data.tuteur_academique_nom:
            p3 = doc.add_paragraph()
            p3.add_run(f"Je remercie également {data.tuteur_academique_nom}")
            if data.tuteur_academique_poste:
                p3.add_run(f", {data.tuteur_academique_poste},")
            p3.add_run(" pour son suivi académique.")

        p4 = doc.add_paragraph()
        run = p4.add_run("[Compléter les remerciements...]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #                       RÉSUMÉ / ABSTRACT
    # ════════════════════════════════════════════════════════════════
    if data.include_abstract:
        resume_heading = doc.add_heading("RÉSUMÉ", level=1)
        resume_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        run = p.add_run("[Résumé du rapport en français...]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        doc.add_heading("Abstract", level=2)
        p2 = doc.add_paragraph()
        run2 = p2.add_run("[English abstract...]")
        run2.italic = True
        run2.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #                          CHAPITRES
    # ════════════════════════════════════════════════════════════════
    for chapter_idx, chapter in enumerate(data.chapters, 1):
        # Numérotation automatique des chapitres
        chapter_num = chapter_idx
        doc.add_heading(f"{chapter_num}. {chapter.title}", level=1)

        p = doc.add_paragraph()
        run = p.add_run(get_chapter_hint(chapter.title))
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        for sub_idx, sub in enumerate(chapter.children, 1):
            # Numérotation automatique des sous-chapitres
            doc.add_heading(f"{chapter_num}.{sub_idx}. {sub.title}", level=2)

            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("[Contenu à rédiger...]")
            run_sub.italic = True
            run_sub.font.color.rgb = RGBColor(128, 128, 128)

            # Gérer les sous-sous-chapitres si présents
            if hasattr(sub, 'children') and sub.children:
                for subsub_idx, subsub in enumerate(sub.children, 1):
                    doc.add_heading(f"{chapter_num}.{sub_idx}.{subsub_idx}. {subsub.title}", level=3)

                    p_subsub = doc.add_paragraph()
                    run_subsub = p_subsub.add_run("[Contenu à rédiger...]")
                    run_subsub.italic = True
                    run_subsub.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════
    #                           ANNEXES
    # ════════════════════════════════════════════════════════════════
    if data.include_annexes:
        annexes_heading = doc.add_heading("ANNEXES", level=1)
        annexes_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("Annexe A - [Titre]", level=2)
        p = doc.add_paragraph()
        run = p.add_run("[Contenu de l'annexe...]")
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    # Sauvegarder
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_chapter_hint(title: str) -> str:
    """Retourne une indication selon le type de chapitre."""
    t = title.lower()
    if "introduction" in t:
        return "[Présenter le contexte, les objectifs et le plan du rapport...]"
    elif "entreprise" in t or "présentation" in t:
        return "[Présenter l'entreprise, son histoire, ses activités, son organisation...]"
    elif "mission" in t:
        return "[Décrire les missions confiées et leurs objectifs...]"
    elif "travail" in t or "réalis" in t:
        return "[Détailler le travail effectué, les méthodes et outils utilisés...]"
    elif "bilan" in t:
        return "[Analyser les résultats, difficultés et compétences acquises...]"
    elif "conclusion" in t:
        return "[Synthétiser les apports du stage et les perspectives...]"
    return "[Contenu à rédiger...]"
