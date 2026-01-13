"""
Générateurs de pages de garde pour le rapport de stage.
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..utils import (
    hex_to_rgb,
    decode_base64_image,
    remove_table_borders,
    add_centered_paragraph,
    add_cell_paragraph,
    set_cell_shading,
    set_table_border,
    set_table_border_top,
)


def generate_cover_classique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style classique - NORMES OFFICIELLES."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Logos en haut
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=3)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(5)
        logo_table.columns[1].width = Cm(6)
        logo_table.columns[2].width = Cm(5)
        remove_table_borders(logo_table)
        row = logo_table.rows[0]

        if has_logo_ecole:
            try:
                para = row.cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = row.cells[2].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    add_centered_paragraph(doc, 25)

    # Titre
    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    # Sujet
    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Image centrale
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # Étudiant
    add_centered_paragraph(doc, 18)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_centered_paragraph(doc, 20)
    run = school.add_run(f"{data.ecole or '[Établissement]'} — {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    # Entreprise
    ent_name = add_centered_paragraph(doc, 4)
    run = ent_name.add_run(f"Stage réalisé chez {data.entreprise_nom or '[Entreprise]'}")
    run.font.size = Pt(14)

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    # Dates
    dates = add_centered_paragraph(doc, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    # Tuteurs
    tuteur_table = doc.add_table(rows=1, cols=2)
    tuteur_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tuteur_table.autofit = False
    tuteur_table.columns[0].width = Cm(7)
    tuteur_table.columns[1].width = Cm(7)
    remove_table_borders(tuteur_table)

    cell_left = tuteur_table.rows[0].cells[0]
    p1 = cell_left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.first_line_indent = Cm(0)
    r1 = p1.add_run("Tuteur entreprise\n")
    r1.font.size = Pt(10)
    run1 = p1.add_run(data.tuteur_nom or "[Nom]")
    run1.bold = True
    run1.font.size = Pt(12)

    cell_right = tuteur_table.rows[0].cells[1]
    p2 = cell_right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.first_line_indent = Cm(0)
    r2 = p2.add_run("Tuteur académique\n")
    r2.font.size = Pt(10)
    run2 = p2.add_run(data.tuteur_academique_nom or "[Nom]")
    run2.bold = True
    run2.font.size = Pt(12)


def generate_cover_moderne(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style moderne avec bandeau coloré."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Bandeau supérieur coloré
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Cm(16)
    remove_table_borders(header_table)
    header_cell = header_table.rows[0].cells[0]
    set_cell_shading(header_cell, data.style.title1_color.lstrip('#'))

    p_header = header_cell.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.first_line_indent = Cm(0)
    p_header.paragraph_format.space_before = Pt(18)
    p_header.paragraph_format.space_after = Pt(6)
    run = p_header.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(255, 255, 255)

    if data.sujet_stage:
        p_sujet = add_cell_paragraph(header_cell, 12)
        p_sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.italic = True

    # Logos
    add_centered_paragraph(doc, 15)
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(7)
        logo_table.columns[1].width = Cm(7)
        remove_table_borders(logo_table)
        row = logo_table.rows[0]

        if has_logo_ecole:
            try:
                para = row.cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = row.cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    # Image centrale
    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    # Étudiant
    add_centered_paragraph(doc, 18)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  •  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    # Entreprise
    ent_name = add_centered_paragraph(doc, 4)
    run = ent_name.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    dates = add_centered_paragraph(doc, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    # Tuteurs
    tuteur_table = doc.add_table(rows=1, cols=2)
    tuteur_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tuteur_table.autofit = False
    tuteur_table.columns[0].width = Cm(7)
    tuteur_table.columns[1].width = Cm(7)
    remove_table_borders(tuteur_table)

    for i, (label, name) in enumerate([
        ("Tuteur entreprise", data.tuteur_nom),
        ("Tuteur académique", data.tuteur_academique_nom)
    ]):
        cell = tuteur_table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r1 = p.add_run(f"{label}\n")
        r1.font.size = Pt(10)
        r2 = p.add_run(name or "[Nom]")
        r2.bold = True
        r2.font.size = Pt(12)


def generate_cover_elegant(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style élégant avec ligne verticale."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Structure : Ligne verticale + Contenu
    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    main_table.columns[0].width = Cm(0.5)
    main_table.columns[1].width = Cm(15.5)

    stripe = main_table.rows[0].cells[0]
    content = main_table.rows[0].cells[1]

    set_cell_shading(stripe, data.style.title1_color.lstrip('#'))
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Logos
    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(2))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(2))
            except:
                pass

    add_cell_paragraph(content, 25)

    title = add_cell_paragraph(content, 4)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 12)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_cell_paragraph(content, 18)
    student = add_cell_paragraph(content, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(content, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_cell_paragraph(content, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  |  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(12)

    ent = add_cell_paragraph(content, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_cell_paragraph(content, 6)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(12)

    dates = add_cell_paragraph(content, 15)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(12)

    tuteurs = add_cell_paragraph(content, 4)
    run = tuteurs.add_run(f"Tuteur entreprise : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(12)

    if data.tuteur_academique_nom:
        tut2 = add_cell_paragraph(content, 0)
        run = tut2.add_run(f"Tuteur académique : {data.tuteur_academique_nom}")
        run.font.size = Pt(12)


def generate_cover_minimaliste(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style minimaliste - Ultra épuré."""
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    for _ in range(5):
        add_centered_paragraph(doc, 18)

    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    line = add_centered_paragraph(doc, 15)
    run = line.add_run("─────────────")
    run.font.size = Pt(12)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 25)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    add_centered_paragraph(doc, 25)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 15)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)

    add_centered_paragraph(doc, 30)

    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    dates = add_centered_paragraph(doc, 4)
    run = dates.add_run(f"{date_debut_fr} — {date_fin_fr}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(128, 128, 128)


def generate_cover_academique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style académique - Cadre double traditionnel."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Cadre extérieur
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_table.autofit = False
    outer_table.columns[0].width = Cm(16)
    set_table_border(outer_table, data.style.title1_color, '24')

    outer_cell = outer_table.rows[0].cells[0]
    add_cell_paragraph(outer_cell, 4)

    # Cadre intérieur
    inner_table = outer_cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inner_table.autofit = False
    inner_table.columns[0].width = Cm(15)
    set_table_border(inner_table, data.style.title1_color, '12')

    content = inner_table.rows[0].cells[0]
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    # Logos
    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    add_cell_paragraph(content, 10)
    school_p = add_cell_paragraph(content, 6)
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school_p.add_run(data.ecole or "[Établissement]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    form_p = add_cell_paragraph(content, 15)
    form_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = form_p.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    title = add_cell_paragraph(content, 6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 12)
        sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 15)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_cell_paragraph(content, 10)
    student = add_cell_paragraph(content, 4)
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run("Présenté par")
    run.font.size = Pt(10)

    name_p = add_cell_paragraph(content, 8)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    ent = add_cell_paragraph(content, 4)
    ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ent.add_run(f"Stage effectué chez {data.entreprise_nom or '[Entreprise]'}")
    run.font.size = Pt(12)

    dates = add_cell_paragraph(content, 15)
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(11)

    tuteurs = add_cell_paragraph(content, 4)
    tuteurs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tuteurs.add_run(f"Tuteur entreprise : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)

    if data.tuteur_academique_nom:
        tut2 = add_cell_paragraph(content, 4)
        tut2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tut2.add_run(f"Tuteur académique : {data.tuteur_academique_nom}")
        run.font.size = Pt(10)

    add_cell_paragraph(content, 10)
    year_p = add_cell_paragraph(content, 6)
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_p.add_run(data.annee_scolaire or "[Année scolaire]")
    run.font.size = Pt(12)
    run.font.color.rgb = primary_color


def generate_cover_geometrique(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style géométrique - Formes modernes."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Bloc coloré en haut à droite
    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    top_table.autofit = False
    top_table.columns[0].width = Cm(10)
    top_table.columns[1].width = Cm(6)
    remove_table_borders(top_table)

    right_cell = top_table.rows[0].cells[1]
    set_cell_shading(right_cell, data.style.title1_color.lstrip('#'))

    left_cell = top_table.rows[0].cells[0]
    if has_logo_ecole:
        try:
            p = left_cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            img = decode_base64_image(data.logos.logo_ecole)
            p.add_run().add_picture(img, height=Cm(2))
        except:
            pass

    p_year = right_cell.paragraphs[0]
    p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_year.paragraph_format.first_line_indent = Cm(0)
    p_year.paragraph_format.space_before = Pt(15)
    p_year.paragraph_format.space_after = Pt(15)
    run = p_year.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.bold = True

    add_centered_paragraph(doc, 20)

    title = doc.add_paragraph()
    title.paragraph_format.first_line_indent = Cm(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = doc.add_paragraph()
        sujet.paragraph_format.first_line_indent = Cm(0)
        sujet.paragraph_format.space_after = Pt(15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    # Ligne décorative
    line_table = doc.add_table(rows=1, cols=1)
    line_table.autofit = False
    line_table.columns[0].width = Cm(5)
    line_cell = line_table.rows[0].cells[0]
    set_cell_shading(line_cell, data.style.title1_color.lstrip('#'))
    remove_table_borders(line_table)
    add_cell_paragraph(line_cell, 0)

    if has_image_centrale:
        try:
            add_centered_paragraph(doc, 15)
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_centered_paragraph(doc, 20)
    student = doc.add_paragraph()
    student.paragraph_format.first_line_indent = Cm(0)
    student.paragraph_format.space_after = Pt(4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = doc.add_paragraph()
    formation.paragraph_format.first_line_indent = Cm(0)
    formation.paragraph_format.space_after = Pt(4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = doc.add_paragraph()
    school.paragraph_format.first_line_indent = Cm(0)
    school.paragraph_format.space_after = Pt(20)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Bloc infos en bas
    info_table = doc.add_table(rows=2, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.autofit = False
    info_table.columns[0].width = Cm(8)
    info_table.columns[1].width = Cm(8)
    remove_table_borders(info_table)

    p1 = info_table.rows[0].cells[0].paragraphs[0]
    p1.paragraph_format.first_line_indent = Cm(0)
    r1 = p1.add_run("Entreprise : ")
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(100, 100, 100)
    r2 = p1.add_run(data.entreprise_nom or "[Entreprise]")
    r2.font.size = Pt(11)
    r2.bold = True

    p2 = info_table.rows[0].cells[1].paragraphs[0]
    p2.paragraph_format.first_line_indent = Cm(0)
    r = p2.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    r.font.size = Pt(10)

    p3 = info_table.rows[1].cells[0].paragraphs[0]
    p3.paragraph_format.first_line_indent = Cm(0)
    r = p3.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    r.font.size = Pt(10)

    if has_logo_entreprise:
        try:
            p4 = info_table.rows[1].cells[1].paragraphs[0]
            p4.paragraph_format.first_line_indent = Cm(0)
            p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            img = decode_base64_image(data.logos.logo_entreprise)
            p4.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass


def generate_cover_bicolore(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style bicolore - Division verticale."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    main_table.columns[0].width = Cm(6.4)
    main_table.columns[1].width = Cm(9.6)

    left_col = main_table.rows[0].cells[0]
    right_col = main_table.rows[0].cells[1]

    set_cell_shading(left_col, data.style.title1_color.lstrip('#'))
    left_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    if has_logo_ecole:
        try:
            add_cell_paragraph(left_col, 15)
            logo_p = add_cell_paragraph(left_col, 15)
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.logo_ecole)
            logo_p.add_run().add_picture(img, height=Cm(2))
        except:
            pass

    for _ in range(3):
        add_cell_paragraph(left_col, 20)

    stage_p = add_cell_paragraph(left_col, 30)
    stage_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stage_p.add_run("STAGE")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(255, 255, 255)

    year_p = add_cell_paragraph(left_col, 10)
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year_p.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(200, 200, 200)

    dates_p = add_cell_paragraph(left_col, 30)
    dates_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates_p.add_run(f"{date_debut_fr}\n—\n{date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(200, 200, 200)

    if has_logo_entreprise:
        try:
            for _ in range(2):
                add_cell_paragraph(left_col, 20)
            logo2_p = add_cell_paragraph(left_col, 10)
            logo2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.logo_entreprise)
            logo2_p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass

    # Colonne droite
    right_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    add_cell_paragraph(right_col, 30)

    title = add_cell_paragraph(right_col, 6)
    run = title.add_run("RAPPORT\nDE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        add_cell_paragraph(right_col, 12)
        sujet = add_cell_paragraph(right_col, 12)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    if has_image_centrale:
        try:
            add_cell_paragraph(right_col, 8)
            img_para = add_cell_paragraph(right_col, 12)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    add_cell_paragraph(right_col, 25)
    student = add_cell_paragraph(right_col, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(right_col, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_cell_paragraph(right_col, 20)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    ent = add_cell_paragraph(right_col, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_cell_paragraph(right_col, 15)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(11)

    add_cell_paragraph(right_col, 20)
    tut = add_cell_paragraph(right_col, 4)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_pro(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style Pro - Corporate business."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Header bar
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    header_table.columns[0].width = Cm(5)
    header_table.columns[1].width = Cm(6)
    header_table.columns[2].width = Cm(5)
    remove_table_borders(header_table)

    for cell in header_table.rows[0].cells:
        set_cell_shading(cell, data.style.title1_color.lstrip('#'))

    row = header_table.rows[0]

    if has_logo_ecole:
        try:
            p = row.cells[0].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            img = decode_base64_image(data.logos.logo_ecole)
            p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass
    else:
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(15)
        p.paragraph_format.space_after = Pt(15)

    p_center = row.cells[1].paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_center.paragraph_format.first_line_indent = Cm(0)
    p_center.paragraph_format.space_before = Pt(15)
    p_center.paragraph_format.space_after = Pt(15)
    run = p_center.add_run(data.annee_scolaire or "[Année]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(255, 255, 255)

    if has_logo_entreprise:
        try:
            p = row.cells[2].paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            img = decode_base64_image(data.logos.logo_entreprise)
            p.add_run().add_picture(img, height=Cm(1.5))
        except:
            pass

    add_centered_paragraph(doc, 30)

    title = add_centered_paragraph(doc, 6)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(14)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_centered_paragraph(doc, 25)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(f"{data.formation or '[Formation]'}  •  {data.ecole or '[Établissement]'}")
    run.font.size = Pt(12)

    add_centered_paragraph(doc, 30)

    # Info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(8)
    remove_table_borders(info_table)

    p1_label = info_table.rows[0].cells[0].paragraphs[0]
    p1_label.paragraph_format.first_line_indent = Cm(0)
    r = p1_label.add_run("Entreprise")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p1_value = info_table.rows[0].cells[1].paragraphs[0]
    p1_value.paragraph_format.first_line_indent = Cm(0)
    r = p1_value.add_run(data.entreprise_nom or "[Entreprise]")
    r.font.size = Pt(12)
    r.bold = True

    p2_label = info_table.rows[1].cells[0].paragraphs[0]
    p2_label.paragraph_format.first_line_indent = Cm(0)
    r = p2_label.add_run("Période")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p2_value = info_table.rows[1].cells[1].paragraphs[0]
    p2_value.paragraph_format.first_line_indent = Cm(0)
    r = p2_value.add_run(f"{date_debut_fr} — {date_fin_fr}")
    r.font.size = Pt(11)

    p3_label = info_table.rows[2].cells[0].paragraphs[0]
    p3_label.paragraph_format.first_line_indent = Cm(0)
    r = p3_label.add_run("Tuteur")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    p3_value = info_table.rows[2].cells[1].paragraphs[0]
    p3_value.paragraph_format.first_line_indent = Cm(0)
    r = p3_value.add_run(data.tuteur_nom or "[Nom]")
    r.font.size = Pt(11)

    if data.tuteur_academique_nom:
        p4_label = info_table.rows[3].cells[0].paragraphs[0]
        p4_label.paragraph_format.first_line_indent = Cm(0)
        r = p4_label.add_run("Suivi")
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)

        p4_value = info_table.rows[3].cells[1].paragraphs[0]
        p4_value.paragraph_format.first_line_indent = Cm(0)
        r = p4_value.add_run(data.tuteur_academique_nom)
        r.font.size = Pt(11)

    # Footer bar
    add_centered_paragraph(doc, 30)
    footer_table = doc.add_table(rows=1, cols=1)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.autofit = False
    footer_table.columns[0].width = Cm(16)

    set_table_border_top(footer_table, data.style.title1_color, '18')
    remove_table_borders(footer_table)

    p_footer = footer_table.rows[0].cells[0].paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.first_line_indent = Cm(0)
    p_footer.paragraph_format.space_before = Pt(10)

    if data.entreprise_ville:
        run = p_footer.add_run(f"{data.entreprise_nom} — {data.entreprise_ville}")
    else:
        run = p_footer.add_run(data.entreprise_nom or "")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_gradient(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style gradient - Dégradé coloré moderne."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    # Bandeau dégradé en haut
    gradient_table = doc.add_table(rows=1, cols=1)
    gradient_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gradient_table.autofit = False
    gradient_table.columns[0].width = Cm(16)
    remove_table_borders(gradient_table)

    grad_cell = gradient_table.rows[0].cells[0]
    set_cell_shading(grad_cell, data.style.title1_color.lstrip('#'))

    p_banner = grad_cell.paragraphs[0]
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.first_line_indent = Cm(0)
    p_banner.paragraph_format.space_before = Pt(20)
    p_banner.paragraph_format.space_after = Pt(8)

    run = p_banner.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(255, 255, 255)

    if data.sujet_stage:
        p_sujet = add_cell_paragraph(grad_cell, 15)
        p_sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.italic = True

    # Logos
    add_centered_paragraph(doc, 15)
    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(7)
        logo_table.columns[1].width = Cm(7)
        remove_table_borders(logo_table)

        if has_logo_ecole:
            try:
                para = logo_table.rows[0].cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = logo_table.rows[0].cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 10)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_centered_paragraph(doc, 20)
    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(12)

    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  •  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    dates = add_centered_paragraph(doc, 10)
    run = dates.add_run(f"Du {date_debut_fr} au {date_fin_fr}")
    run.font.size = Pt(11)

    tut = add_centered_paragraph(doc, 4)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_timeline(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style timeline - Frise chronologique verticale."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.autofit = False
    remove_table_borders(main_table)

    main_table.columns[0].width = Cm(2)
    main_table.columns[1].width = Cm(14)

    timeline_col = main_table.rows[0].cells[0]
    content_col = main_table.rows[0].cells[1]

    # Timeline
    timeline_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    p1 = add_cell_paragraph(timeline_col, 4)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("●")
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    date_p1 = add_cell_paragraph(timeline_col, 20)
    date_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p1.add_run(date_debut_fr.split()[0] if date_debut_fr else "")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(6):
        line_p = add_cell_paragraph(timeline_col, 0)
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line_p.add_run("│")
        run.font.size = Pt(10)
        run.font.color.rgb = primary_color

    add_cell_paragraph(timeline_col, 20)
    p2 = add_cell_paragraph(timeline_col, 4)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("●")
    run.font.size = Pt(14)
    run.font.color.rgb = primary_color

    date_p2 = add_cell_paragraph(timeline_col, 0)
    date_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p2.add_run(date_fin_fr.split()[0] if date_fin_fr else "")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Content
    content_col.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    if has_logo_ecole or has_logo_entreprise:
        logo_tbl = content_col.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(6)
        logo_tbl.columns[1].width = Cm(6)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

    add_cell_paragraph(content_col, 15)
    title = add_cell_paragraph(content_col, 4)
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_cell_paragraph(content_col, 10)
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content_col, 10)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    add_cell_paragraph(content_col, 15)
    student = add_cell_paragraph(content_col, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(content_col, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)

    school = add_cell_paragraph(content_col, 10)
    run = school.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    ent = add_cell_paragraph(content_col, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    add_cell_paragraph(content_col, 10)
    tut = add_cell_paragraph(content_col, 0)
    run = tut.add_run(f"Tuteur : {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)


def generate_cover_creative(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style creative - Design original avec cercles."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100
    primary_color = hex_to_rgb(data.style.title1_color)

    if has_logo_ecole or has_logo_entreprise:
        logo_table = doc.add_table(rows=1, cols=2)
        logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        logo_table.autofit = False
        logo_table.columns[0].width = Cm(8)
        logo_table.columns[1].width = Cm(8)
        remove_table_borders(logo_table)

        if has_logo_ecole:
            try:
                para = logo_table.rows[0].cells[0].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_ecole)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

        if has_logo_entreprise:
            try:
                para = logo_table.rows[0].cells[1].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.first_line_indent = Cm(0)
                img = decode_base64_image(data.logos.logo_entreprise)
                para.add_run().add_picture(img, height=Cm(1.8))
            except:
                pass

    add_centered_paragraph(doc, 30)

    rapport = add_centered_paragraph(doc, 0)
    run = rapport.add_run("RAPPORT")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = primary_color

    de_stage = add_centered_paragraph(doc, 8)
    run = de_stage.add_run("DE STAGE")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(150, 150, 150)

    line = add_centered_paragraph(doc, 15)
    run = line.add_run("━━━━━━━━━━━━━━━━")
    run.font.size = Pt(10)
    run.font.color.rgb = primary_color

    if data.sujet_stage:
        sujet = add_centered_paragraph(doc, 15)
        run = sujet.add_run(f"« {data.sujet_stage} »")
        run.font.size = Pt(13)
        run.italic = True

    if has_image_centrale:
        try:
            img_para = add_centered_paragraph(doc, 15)
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(3))
        except:
            pass

    add_centered_paragraph(doc, 20)

    student = add_centered_paragraph(doc, 4)
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_centered_paragraph(doc, 4)
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    school = add_centered_paragraph(doc, 15)
    run = school.add_run(f"{data.ecole or '[Établissement]'}  |  {data.annee_scolaire or '[Année]'}")
    run.font.size = Pt(10)

    sep = add_centered_paragraph(doc, 10)
    run = sep.add_run("─────────────")
    run.font.color.rgb = RGBColor(200, 200, 200)

    ent = add_centered_paragraph(doc, 4)
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = primary_color

    if data.entreprise_ville:
        ville = add_centered_paragraph(doc, 4)
        run = ville.add_run(data.entreprise_ville)
        run.font.size = Pt(10)

    dates = add_centered_paragraph(doc, 8)
    run = dates.add_run(f"{date_debut_fr}  —  {date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    tut = add_centered_paragraph(doc, 4)
    run = tut.add_run(f"Encadré par {data.tuteur_nom or '[Nom]'}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)


def generate_cover_luxe(doc, data, date_debut_fr, date_fin_fr, duree):
    """Page de garde style luxe - Élégant avec bordures dorées."""
    has_logo_ecole = data.logos.logo_ecole and len(data.logos.logo_ecole) > 100
    has_logo_entreprise = data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100
    has_image_centrale = data.logos.image_centrale and len(data.logos.image_centrale) > 100

    gold_color = RGBColor(184, 134, 11)

    # Cadre extérieur doré
    outer_table = doc.add_table(rows=1, cols=1)
    outer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer_table.autofit = False
    outer_table.columns[0].width = Cm(16)
    set_table_border(outer_table, 'b8860b', '36')

    outer_cell = outer_table.rows[0].cells[0]

    add_cell_paragraph(outer_cell, 6)
    inner_table = outer_cell.add_table(rows=1, cols=1)
    inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inner_table.autofit = False
    inner_table.columns[0].width = Cm(15)
    set_table_border(inner_table, 'b8860b', '18')

    content = inner_table.rows[0].cells[0]
    content.paragraphs[0].paragraph_format.first_line_indent = Cm(0)

    if has_logo_ecole or has_logo_entreprise:
        add_cell_paragraph(content, 6)
        logo_tbl = content.add_table(rows=1, cols=2)
        logo_tbl.autofit = False
        logo_tbl.columns[0].width = Cm(7)
        logo_tbl.columns[1].width = Cm(7)
        remove_table_borders(logo_tbl)

        if has_logo_ecole:
            try:
                p = logo_tbl.rows[0].cells[0].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_ecole)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

        if has_logo_entreprise:
            try:
                p = logo_tbl.rows[0].cells[1].paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img = decode_base64_image(data.logos.logo_entreprise)
                p.add_run().add_picture(img, height=Cm(1.5))
            except:
                pass

    add_cell_paragraph(content, 10)
    school_p = add_cell_paragraph(content, 4)
    school_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = school_p.add_run(data.ecole or "[Établissement]")
    run.font.size = Pt(11)
    run.font.color.rgb = gold_color
    run.font.small_caps = True

    orn1 = add_cell_paragraph(content, 6)
    orn1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = orn1.add_run("— ✦ —")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color

    title = add_cell_paragraph(content, 6)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAPPORT DE STAGE")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(50, 50, 50)

    if data.sujet_stage:
        sujet = add_cell_paragraph(content, 8)
        sujet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sujet.add_run(data.sujet_stage)
        run.font.size = Pt(13)
        run.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)

    if has_image_centrale:
        try:
            img_para = add_cell_paragraph(content, 10)
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img = decode_base64_image(data.logos.image_centrale)
            img_para.add_run().add_picture(img, height=Cm(2.5))
        except:
            pass

    add_cell_paragraph(content, 10)
    line = add_cell_paragraph(content, 8)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("───────────────────")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color

    student = add_cell_paragraph(content, 6)
    student.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = student.add_run(f"{data.prenom or '[Prénom]'} {data.nom or '[NOM]'}")
    run.bold = True
    run.font.size = Pt(16)

    formation = add_cell_paragraph(content, 2)
    formation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formation.add_run(data.formation or "[Formation]")
    run.font.size = Pt(11)
    run.italic = True

    add_cell_paragraph(content, 10)
    ent = add_cell_paragraph(content, 2)
    ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ent.add_run(data.entreprise_nom or "[Entreprise]")
    run.font.size = Pt(13)
    run.font.color.rgb = gold_color

    dates = add_cell_paragraph(content, 4)
    dates.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dates.add_run(f"{date_debut_fr}  —  {date_fin_fr}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    add_cell_paragraph(content, 8)
    year = add_cell_paragraph(content, 6)
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run(data.annee_scolaire or "[Année scolaire]")
    run.font.size = Pt(10)
    run.font.color.rgb = gold_color
    run.font.small_caps = True


# Registry des générateurs de covers
COVER_GENERATORS = {
    'classique': generate_cover_classique,
    'moderne': generate_cover_moderne,
    'elegant': generate_cover_elegant,
    'minimaliste': generate_cover_minimaliste,
    'academique': generate_cover_academique,
    'geometrique': generate_cover_geometrique,
    'bicolore': generate_cover_bicolore,
    'pro': generate_cover_pro,
    'gradient': generate_cover_gradient,
    'timeline': generate_cover_timeline,
    'creative': generate_cover_creative,
    'luxe': generate_cover_luxe,
}


def get_cover_generator(model_name: str):
    """Retourne le générateur de cover correspondant au nom du modèle."""
    return COVER_GENERATORS.get(model_name, generate_cover_classique)
