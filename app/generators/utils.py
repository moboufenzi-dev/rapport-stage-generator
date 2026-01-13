"""
Fonctions utilitaires pour la génération de documents Word.
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import base64
from datetime import datetime
from PIL import Image as PILImage


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convertit une couleur hexadécimale en RGBColor."""
    hex_color = hex_color.lstrip('#')
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def decode_base64_image(base64_str: str) -> io.BytesIO:
    """Décode une image base64 et la convertit en PNG pour compatibilité python-docx."""
    if not base64_str:
        raise ValueError("Image base64 vide")
    if ',' in base64_str:
        base64_str = base64_str.split(',')[1]
    decoded = base64.b64decode(base64_str)

    try:
        input_stream = io.BytesIO(decoded)
        pil_image = PILImage.open(input_stream)

        # Convertir en RGB si nécessaire (pour les images RGBA ou autres modes)
        if pil_image.mode in ('RGBA', 'LA', 'P'):
            background = PILImage.new('RGB', pil_image.size, (255, 255, 255))
            if pil_image.mode == 'P':
                pil_image = pil_image.convert('RGBA')
            background.paste(pil_image, mask=pil_image.split()[-1] if pil_image.mode == 'RGBA' else None)
            pil_image = background
        elif pil_image.mode != 'RGB':
            pil_image = pil_image.convert('RGB')

        output_stream = io.BytesIO()
        pil_image.save(output_stream, format='PNG')
        output_stream.seek(0)
        return output_stream
    except Exception as e:
        print(f"[DEBUG] Conversion Pillow échouée: {e}, utilisation directe")
        return io.BytesIO(decoded)


def format_date_fr(date_str: str) -> str:
    """Formate une date en français (ex: 15 janvier 2024)."""
    if not date_str:
        return "[Date]"
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        mois = ["janvier", "février", "mars", "avril", "mai", "juin",
                "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
        return f"{d.day} {mois[d.month-1]} {d.year}"
    except:
        return date_str


def calculate_duration(date_debut: str, date_fin: str) -> str:
    """Calcule la durée entre deux dates en mois."""
    if not date_debut or not date_fin:
        return "[durée]"
    try:
        d1 = datetime.strptime(date_debut, "%Y-%m-%d")
        d2 = datetime.strptime(date_fin, "%Y-%m-%d")
        months = (d2.year - d1.year) * 12 + d2.month - d1.month
        return f"{months} mois" if months != 1 else "1 mois"
    except:
        return "[durée]"


def add_page_number_field(paragraph):
    """Ajoute un champ PAGE pour le numéro de page."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE "

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    text_elem = OxmlElement('w:t')
    text_elem.text = "1"

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(text_elem)
    run._r.append(fldChar3)


def set_cell_shading(cell, color_hex: str):
    """Applique une couleur de fond à une cellule."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def remove_table_borders(table):
    """Supprime toutes les bordures d'un tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def add_centered_paragraph(doc, space_after: int = 0):
    """Crée un paragraphe centré sans indentation (pour page de garde)."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    return para


def create_toc_entry(doc, text: str, level: int = 1, page: str = ""):
    """Crée une entrée de table des matières avec tabulation et points de suite."""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(4)

    # Indentation selon le niveau
    if level == 1:
        paragraph.paragraph_format.left_indent = Cm(0)
    elif level == 2:
        paragraph.paragraph_format.left_indent = Cm(0.75)
    else:
        paragraph.paragraph_format.left_indent = Cm(1.5)

    # Tabulation avec points de suite à 15cm
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    # Texte du titre
    run = paragraph.add_run(text)
    if level == 1:
        run.bold = True
        run.font.size = Pt(11)
    elif level == 2:
        run.font.size = Pt(10)
    else:
        run.font.size = Pt(10)
        run.italic = True

    # Tabulation + numéro de page
    paragraph.add_run("\t")
    page_run = paragraph.add_run(page)
    page_run.font.size = Pt(11 if level == 1 else 10)
    if level == 1:
        page_run.bold = True


def create_toc(doc, data):
    """Crée une table des matières avec numérotation automatique."""
    page_num = 3

    if data.include_thanks:
        create_toc_entry(doc, "REMERCIEMENTS", 1, str(page_num))
        page_num += 1

    if data.include_abstract:
        create_toc_entry(doc, "RÉSUMÉ", 1, str(page_num))
        page_num += 1

    for chapter_idx, chapter in enumerate(data.chapters, 1):
        create_toc_entry(doc, f"{chapter_idx}. {chapter.title}", 1, str(page_num))
        for sub_idx, sub in enumerate(chapter.children, 1):
            create_toc_entry(doc, f"{chapter_idx}.{sub_idx}. {sub.title}", 2, str(page_num))
            if hasattr(sub, 'children') and sub.children:
                for subsub_idx, subsub in enumerate(sub.children, 1):
                    create_toc_entry(doc, f"{chapter_idx}.{sub_idx}.{subsub_idx}. {subsub.title}", 3, str(page_num))
        page_num += 1

    if data.include_annexes:
        create_toc_entry(doc, "ANNEXES", 1, str(page_num))


def setup_document_styles(doc, style_config):
    """Configure les styles du document avec indentation professionnelle."""
    styles = doc.styles

    # Normal - texte avec indentation
    normal = styles['Normal']
    normal.font.name = style_config.font_family
    normal.font.size = Pt(style_config.font_size)
    normal.paragraph_format.line_spacing = style_config.line_spacing
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1 - Chapitres principaux
    h1 = styles['Heading 1']
    h1.font.name = style_config.font_family
    h1.font.size = Pt(style_config.title1_size)
    h1.font.bold = style_config.title1_bold
    h1.font.color.rgb = hex_to_rgb(style_config.title1_color)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.left_indent = Cm(0)
    h1.paragraph_format.first_line_indent = Cm(0)

    # Heading 2 - Sous-sections
    h2 = styles['Heading 2']
    h2.font.name = style_config.font_family
    h2.font.size = Pt(style_config.title2_size)
    h2.font.bold = style_config.title2_bold
    h2.font.color.rgb = hex_to_rgb(style_config.title2_color)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.left_indent = Cm(0.75)
    h2.paragraph_format.first_line_indent = Cm(0)

    # Heading 3 - Sous-sous-sections
    h3 = styles['Heading 3']
    h3.font.name = style_config.font_family
    h3.font.size = Pt(style_config.title3_size)
    h3.font.italic = style_config.title3_italic
    h3.font.bold = False
    h3.font.color.rgb = hex_to_rgb(style_config.title3_color)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.left_indent = Cm(1.5)
    h3.paragraph_format.first_line_indent = Cm(0)


def setup_header_with_logos(section, data):
    """Configure l'en-tête avec logos - symétrique."""
    header = section.header

    # Supprimer le contenu existant
    for para in header.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Créer un tableau symétrique (largeurs fixes) - total 16cm
    tbl = header.add_table(rows=1, cols=3, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    # Largeurs symétriques
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(6)
    tbl.columns[2].width = Cm(5)

    remove_table_borders(tbl)

    row = tbl.rows[0]

    # Logo école (gauche)
    cell_left = row.cells[0]
    para_left = cell_left.paragraphs[0]
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_left.paragraph_format.first_line_indent = Cm(0)
    if data.logos.logo_ecole and len(data.logos.logo_ecole) > 100:
        try:
            img = decode_base64_image(data.logos.logo_ecole)
            run = para_left.add_run()
            run.add_picture(img, height=Cm(1.2))
        except:
            pass

    # Centre (vide)
    cell_center = row.cells[1]
    para_center = cell_center.paragraphs[0]
    para_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_center.paragraph_format.first_line_indent = Cm(0)

    # Logo entreprise (droite)
    cell_right = row.cells[2]
    para_right = cell_right.paragraphs[0]
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_right.paragraph_format.first_line_indent = Cm(0)
    if data.logos.logo_entreprise and len(data.logos.logo_entreprise) > 100:
        try:
            img = decode_base64_image(data.logos.logo_entreprise)
            run = para_right.add_run()
            run.add_picture(img, height=Cm(1.2))
        except:
            pass


def add_cell_paragraph(cell, space_after: int = 0):
    """Crée un paragraphe sans indentation dans une cellule."""
    para = cell.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.space_after = Pt(space_after)
    return para


def set_table_border(table, color_hex: str, size: str = '6'):
    """Applique une bordure colorée à un tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color_hex.lstrip('#'))
        tbl_borders.append(border)
    # Intérieurs vides
    for border_name in ['insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def set_table_border_top(table, color_hex: str, size: str = '6'):
    """Applique une bordure uniquement en haut du tableau."""
    tbl_element = table._tbl
    tbl_pr = tbl_element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl_element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    border = OxmlElement('w:top')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), size)
    border.set(qn('w:color'), color_hex.lstrip('#'))
    tbl_borders.append(border)
    tbl_pr.append(tbl_borders)


def setup_footer_with_page_number(section, data):
    """Configure le pied de page : entreprise à gauche, numéro au centre, nom à droite."""
    footer = section.footer

    # Supprimer le contenu existant
    for para in footer.paragraphs:
        p = para._element
        p.getparent().remove(p)

    # Créer un tableau symétrique pour le footer - total 16cm
    tbl = footer.add_table(rows=1, cols=3, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(6)
    tbl.columns[2].width = Cm(5)

    remove_table_borders(tbl)

    row = tbl.rows[0]

    # Entreprise (gauche)
    cell_left = row.cells[0]
    para_left = cell_left.paragraphs[0]
    para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_left.paragraph_format.first_line_indent = Cm(0)
    if data.entreprise_nom:
        run = para_left.add_run(data.entreprise_nom)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)

    # Numéro de page (centre)
    cell_center = row.cells[1]
    para_center = cell_center.paragraphs[0]
    para_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_center.paragraph_format.first_line_indent = Cm(0)
    if data.page.show_page_number:
        run = para_center.add_run("- ")
        run.font.size = Pt(9)
        add_page_number_field(para_center)
        run = para_center.add_run(" -")
        run.font.size = Pt(9)

    # Nom étudiant (droite)
    cell_right = row.cells[2]
    para_right = cell_right.paragraphs[0]
    para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_right.paragraph_format.first_line_indent = Cm(0)
    if data.page.show_student_name and data.nom:
        run = para_right.add_run(f"{data.prenom} {data.nom}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
