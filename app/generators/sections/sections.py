"""
Générateurs de sections du rapport de stage.
"""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..utils import create_toc


def generate_toc_section(doc, data):
    """Génère la table des matières."""
    toc_heading = doc.add_heading("TABLE DES MATIÈRES", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    create_toc(doc, data)

    doc.add_page_break()


def generate_thanks_section(doc, data):
    """Génère la section remerciements."""
    thanks_heading = doc.add_heading("REMERCIEMENTS", level=1)
    thanks_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = doc.add_paragraph()
    p1.add_run(f"Je tiens à remercier {data.entreprise_nom or '[Entreprise]'} pour m'avoir accueilli durant ce stage.")

    if data.tuteur_nom:
        p2 = doc.add_paragraph()
        p2.add_run(f"Je remercie particulièrement {data.tuteur_nom}")
        if data.tuteur_poste:
            p2.add_run(f", {data.tuteur_poste},")
        p2.add_run(" pour son encadrement tout au long de ce stage.")

    if data.tuteur_academique_nom:
        p3 = doc.add_paragraph()
        p3.add_run(f"Je remercie également {data.tuteur_academique_nom}")
        if data.tuteur_academique_poste:
            p3.add_run(f", {data.tuteur_academique_poste},")
        p3.add_run(" pour son suivi académique.")

    p4 = doc.add_paragraph()
    run = p4.add_run("[Compléter les remerciements...]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def generate_abstract_section(doc, data):
    """Génère la section résumé/abstract."""
    resume_heading = doc.add_heading("RÉSUMÉ", level=1)
    resume_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    run = p.add_run("[Résumé du rapport en français...]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    doc.add_heading("Abstract", level=2)
    p2 = doc.add_paragraph()
    run2 = p2.add_run("[English abstract...]")
    run2.italic = True
    run2.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()


def get_chapter_hint(title: str) -> str:
    """Retourne une indication selon le type de chapitre."""
    t = title.lower()
    if "introduction" in t:
        return "[Présenter le contexte, les objectifs et le plan du rapport...]"
    elif "entreprise" in t or "présentation" in t:
        return "[Présenter l'entreprise, son histoire, ses activités, son organisation...]"
    elif "mission" in t:
        return "[Décrire les missions confiées et leurs objectifs...]"
    elif "travail" in t or "réalis" in t:
        return "[Détailler le travail effectué, les méthodes et outils utilisés...]"
    elif "bilan" in t:
        return "[Analyser les résultats, difficultés et compétences acquises...]"
    elif "conclusion" in t:
        return "[Synthétiser les apports du stage et les perspectives...]"
    return "[Contenu à rédiger...]"


def generate_chapters(doc, data):
    """Génère tous les chapitres du rapport."""
    for chapter_idx, chapter in enumerate(data.chapters, 1):
        chapter_num = chapter_idx
        doc.add_heading(f"{chapter_num}. {chapter.title}", level=1)

        p = doc.add_paragraph()
        run = p.add_run(get_chapter_hint(chapter.title))
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        for sub_idx, sub in enumerate(chapter.children, 1):
            doc.add_heading(f"{chapter_num}.{sub_idx}. {sub.title}", level=2)

            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("[Contenu à rédiger...]")
            run_sub.italic = True
            run_sub.font.color.rgb = RGBColor(128, 128, 128)

            if hasattr(sub, 'children') and sub.children:
                for subsub_idx, subsub in enumerate(sub.children, 1):
                    doc.add_heading(f"{chapter_num}.{sub_idx}.{subsub_idx}. {subsub.title}", level=3)

                    p_subsub = doc.add_paragraph()
                    run_subsub = p_subsub.add_run("[Contenu à rédiger...]")
                    run_subsub.italic = True
                    run_subsub.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_page_break()


def generate_annexes_section(doc, data):
    """Génère la section annexes."""
    annexes_heading = doc.add_heading("ANNEXES", level=1)
    annexes_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Annexe A - [Titre]", level=2)
    p = doc.add_paragraph()
    run = p.add_run("[Contenu de l'annexe...]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
